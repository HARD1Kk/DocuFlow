"""
Tests for DocumentLoader - Data Source Layer
"""

import pytest
from docx import Document

from docuflow.core.loaders import DocumentLoader
from docuflow.schemas import RawDocument


@pytest.fixture
def loader():
    """Create a DocumentLoader instance"""
    return DocumentLoader()


@pytest.fixture
def txt_file(tmp_path):
    """Create a temporary TXT file"""
    file = tmp_path / "sample_test.txt"
    file.write_text("IntroduCtioN to TXT\n\nThis is a text file.\n\nDetails\n\nSome content.")

    return file


@pytest.fixture
def md_file(tmp_path):
    """Create a temporary Markdown file"""
    file = tmp_path / "sample_test.md"
    file.write_text("# Introduction to MD\n\nThis is markdown.\n\n## Details\n\nMore content.")
    return file


@pytest.fixture
def docx_file(tmp_path):
    """Create a temporary DOCX file"""
    file = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is a DOCX file.")
    doc.add_heading("Details", level=2)
    doc.add_paragraph("Some details here.")
    doc.save(str(file))
    return file


class TestDocumentLoaderValidation:
    """Tests for file validation"""

    def test_validation_existing_files(self, loader, txt_file):
        """Test validation passes for existing TXT file"""
        assert loader.validate(str(txt_file)) is True

    def test_validate_existing_md_file(self, loader, md_file):
        """Test validation passes for existing MD file"""
        assert loader.validate(str(md_file)) is True

    def test_validate_nonexistent_file(self, loader):
        """Test validation fails for missing files"""
        assert loader.validate("/nonexistent/file.txt") is False

    def test_validate_unsupported_format(self, loader, tmp_path):
        """Test validation fails for unsupported formats"""
        unsupported = tmp_path / "file.xyz"
        unsupported.write_text("test")

        assert loader.validate(str(unsupported)) is False

    def test_supported_extensions(self, loader):
        """Test loader declares all supported formats"""
        assert ".pdf" in loader.SUPPORTED_EXTENSIONS
        assert ".docx" in loader.SUPPORTED_EXTENSIONS
        assert ".txt" in loader.SUPPORTED_EXTENSIONS
        assert ".md" in loader.SUPPORTED_EXTENSIONS


class TestDocumentLoaderLoading:
    """Test document loading"""

    def test_load_txt_file(self, loader, txt_file):
        """Test loading a TXT file returns RawDocument"""
        result = loader.load(str(txt_file))
        # Check result is list with one item
        assert isinstance(result, list)

        assert len(result) == 1

        # Check it's a RawDocument
        raw_doc = result[0]

        assert isinstance(raw_doc, RawDocument)

        # Check content
        assert raw_doc.content is not None
        assert len(raw_doc.content) > 0
        assert "introduction" in raw_doc.content.lower()

    def test_load_md_file(self, loader, md_file):
        """Test loading a MD file returns RawDocument"""
        result = loader.load(str(md_file))

        assert isinstance(result, list)

        assert len(result) == 1
        raw_doc = result[0]

        assert isinstance(raw_doc, RawDocument)
        assert "# introduction to md" in raw_doc.content.lower()

    def test_load_docx_file(self, loader, docx_file):
        """Test loading a DOCX file (converts to markdown)"""

        result = loader.load(str(docx_file))
        assert len(result) == 1
        raw_doc = result[0]

        assert isinstance(raw_doc, RawDocument)
        assert raw_doc.content is not None
        # DOCX converted to markdown should have some content
        assert len(raw_doc.content) > 0
